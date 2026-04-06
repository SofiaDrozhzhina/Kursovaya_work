"""
docs_generator.py — генерация документов для студентов

Документы:
  1. Согласие на обработку персональных данных (.docx)
  2. Справка об обучении (.docx)
  3. Сертификат об окончании курса (.docx)
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io
from datetime import datetime


# ─── Helpers ─────────────────────────────────────────────────────────────────

def _set_font(run, name="Times New Roman", size=12, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    # Fix for Cyrillic fonts
    run._element.rPr.rFonts.set(qn('w:cs'), name)


def _para(doc, text="", align=WD_ALIGN_PARAGRAPH.LEFT, size=12,
          bold=False, italic=False, space_before=0, space_after=6,
          font="Times New Roman", color=None):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        run = p.add_run(text)
        _set_font(run, font, size, bold, italic, color)
    return p


def _add_run(para, text, size=12, bold=False, italic=False,
             font="Times New Roman", color=None):
    run = para.add_run(text)
    _set_font(run, font, size, bold, italic, color)
    return run


def _set_page_margins(doc, top=2, bottom=2, left=3, right=1.5):
    """Стандартные поля для документов РФ (в см)"""
    section = doc.sections[0]
    section.top_margin    = Cm(top)
    section.bottom_margin = Cm(bottom)
    section.left_margin   = Cm(left)
    section.right_margin  = Cm(right)


def _add_line(doc):
    """Горизонтальная линия"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '4B3F8C')
    pBdr.append(bottom)
    pPr.append(pBdr)


def _signature_row(doc, left_label, right_label=""):
    """Строка с подписью: слева должность, справа дата/подпись"""
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    # убрать границы
    for cell in table.rows[0].cells:
        for border in ('top','bottom','left','right'):
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            bd = OxmlElement(f'w:{border}')
            bd.set(qn('w:val'), 'nil')
            tcBorders.append(bd)
            tcPr.append(tcBorders)

    left_cell  = table.rows[0].cells[0]
    right_cell = table.rows[0].cells[1]
    left_cell.width  = Cm(10)
    right_cell.width = Cm(6)

    lp = left_cell.paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    lr = lp.add_run(left_label)
    _set_font(lr, size=11)

    rp = right_cell.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rr = rp.add_run(right_label)
    _set_font(rr, size=11)
    return table


def _to_buffer(doc) -> bytes:
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ─── 1. Согласие на обработку персональных данных ────────────────────────────

def generate_consent(student_name: str, student_email: str = "",
                     group_name: str = "") -> bytes:
    """Согласие на обработку персональных данных"""
    doc = Document()
    _set_page_margins(doc)
    now = datetime.now()

    # Шапка
    _para(doc, "ГОСУДАРСТВЕННОЕ БЮДЖЕТНОЕ ОБРАЗОВАТЕЛЬНОЕ УЧРЕЖДЕНИЕ",
          WD_ALIGN_PARAGRAPH.CENTER, size=12, bold=True, space_after=0)
    _para(doc, "ВЫСШЕГО ПРОФЕССИОНАЛЬНОГО ОБРАЗОВАНИЯ",
          WD_ALIGN_PARAGRAPH.CENTER, size=12, bold=True, space_after=0)
    _para(doc, "«УНИВЕРСИТЕТ»",
          WD_ALIGN_PARAGRAPH.CENTER, size=13, bold=True, space_after=12)

    _add_line(doc)

    # Заголовок
    _para(doc, "СОГЛАСИЕ",
          WD_ALIGN_PARAGRAPH.CENTER, size=16, bold=True,
          space_before=12, space_after=4)
    _para(doc, "на обработку персональных данных",
          WD_ALIGN_PARAGRAPH.CENTER, size=13, bold=True, space_after=16)

    # Дата и место
    months_ru = {
        'January':'января','February':'февраля','March':'марта','April':'апреля',
        'May':'мая','June':'июня','July':'июля','August':'августа',
        'September':'сентября','October':'октября','November':'ноября','December':'декабря'
    }
    day_str  = now.strftime('%d')
    mon_str  = months_ru[now.strftime('%B')]
    year_str = now.strftime('%Y')
    date_full = f'«{day_str}» {mon_str} {year_str} г.'
    date_str  = f'{day_str} {mon_str} {year_str}'

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(12)
    _add_run(p,
             f"г. Москва                                                         {date_full}",
             size=11, italic=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.space_after = Pt(8)
    _add_run(p, "Я, ")
    _add_run(p, student_name, bold=True)
    _add_run(p,
             f", студент{'ка' if student_name.endswith('а') or student_name.endswith('я') else ''} "
             f"{'группы ' + group_name + ',' if group_name else ''} "
             "в соответствии с Федеральным законом от 27.07.2006 № 152-ФЗ "
             "«О персональных данных», свободно, своей волей и в своём интересе "
             "даю согласие Университету (далее — Оператор) на обработку "
             "следующих моих персональных данных:")

    # Список данных
    items = [
        "фамилия, имя, отчество;",
        "дата и место рождения;",
        "паспортные данные;",
        "адрес регистрации и фактического проживания;",
        "контактный номер телефона;",
        f"адрес электронной почты{': ' + student_email if student_email else ';'}",
        "сведения об обучении (группа, курс, успеваемость).",
    ]
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.left_indent  = Cm(1.5)
        p.paragraph_format.space_after  = Pt(3)
        run = p.add_run(item)
        _set_font(run, size=12)

    doc.add_paragraph()

    for text in [
        "Обработка персональных данных осуществляется в целях обеспечения "
        "учебного процесса, формирования личного дела студента, выдачи "
        "документов об образовании и иных целях, непосредственно связанных "
        "с образовательной деятельностью Оператора.",

        "Согласие предоставляется на совершение следующих действий с "
        "персональными данными: сбор, запись, систематизация, накопление, "
        "хранение, уточнение (обновление, изменение), извлечение, "
        "использование, передача (распространение, предоставление, доступ), "
        "обезличивание, блокирование, удаление, уничтожение.",

        "Согласие действует с момента подписания и до его отзыва. "
        "Я вправе отозвать настоящее согласие, направив письменное "
        "заявление в адрес Оператора.",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = Cm(1.25)
        p.paragraph_format.space_after = Pt(8)
        _add_run(p, text, size=12)

    doc.add_paragraph()
    _add_line(doc)
    doc.add_paragraph()

    # Подписи
    _signature_row(doc,
                   f"Студент: {student_name}",
                   f"Дата: {now.strftime('%d.%m.%Y')}")
    doc.add_paragraph()
    _signature_row(doc,
                   "Подпись: ____________________",
                   "")
    doc.add_paragraph()
    _para(doc,
          "Ректор Университета: ____________________  / И.О. Фамилия /",
          size=11, space_before=8)

    return _to_buffer(doc)


# ─── 2. Справка об обучении ───────────────────────────────────────────────────

def generate_study_certificate(student_name: str, group_name: str,
                               enrollment_year: int,
                               courses: list) -> bytes:
    """
    Справка об обучении.
    courses — список dict: {title, status, grade}
    """
    doc = Document()
    _set_page_margins(doc)
    now = datetime.now()

    months_ru = {
        'January':'января','February':'февраля','March':'марта','April':'апреля',
        'May':'мая','June':'июня','July':'июля','August':'августа',
        'September':'сентября','October':'октября','November':'ноября','December':'декабря'
    }
    _day  = now.strftime('%d')
    _mon  = months_ru[now.strftime('%B')]
    _year = now.strftime('%Y')

    # Угловой штамп (имитация)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    _add_run(p, "Государственное бюджетное образовательное\nучреждение «Университет»",
             size=10, italic=True)

    _para(doc, f"Исх. № ____  от {_day}.{_mon}.{_year}",
          size=10, italic=True, space_after=20)

    # Заголовок
    _para(doc, "С П Р А В К А",
          WD_ALIGN_PARAGRAPH.CENTER, size=16, bold=True, space_after=4)
    _para(doc, "об обучении",
          WD_ALIGN_PARAGRAPH.CENTER, size=13, italic=True, space_after=20)

    _add_line(doc)
    doc.add_paragraph()

    # Основной текст
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(1.25)
    _add_run(p, "Настоящая справка выдана в том, что ")
    _add_run(p, student_name, bold=True)
    status_text = (
        f" действительно является студентом{'кой' if student_name.split()[-1].endswith(('а','я','ая','яя')) else 'ом'} "
        f"{'группы ' + group_name if group_name else 'Университета'}, "
        f"{'поступившим' if not student_name.split()[-1].endswith(('а','я')) else 'поступившей'} "
        f"в {enrollment_year} году."
    )
    _add_run(p, status_text)

    doc.add_paragraph()

    # Таблица курсов
    _para(doc, "Сведения о пройденных курсах:", size=12, bold=True, space_after=6)

    if courses:
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        widths = [Cm(1.2), Cm(7.5), Cm(3.5), Cm(3.0)]

        # Шапка таблицы
        hdr_cells = table.rows[0].cells
        headers = ["№", "Наименование курса", "Статус", "Оценка"]
        for i, (cell, hdr, w) in enumerate(zip(hdr_cells, headers, widths)):
            cell.width = w
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(hdr)
            _set_font(run, size=11, bold=True)
            # Заливка шапки
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'),   'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'),  'E8E4F8')
            tcPr.append(shd)

        status_map = {'completed': 'Завершён', 'active': 'В процессе', 'dropped': 'Отчислен'}
        for idx, course in enumerate(courses, 1):
            row_cells = table.add_row().cells
            vals = [
                str(idx),
                course.get('title', ''),
                status_map.get(course.get('status',''), course.get('status','')),
                str(course.get('grade', '—')) if course.get('grade') is not None else '—',
            ]
            aligns = [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT,
                      WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER]
            for cell, val, align, w in zip(row_cells, vals, aligns, widths):
                cell.width = w
                p = cell.paragraphs[0]
                p.alignment = align
                run = p.add_run(val)
                _set_font(run, size=11)
    else:
        _para(doc, "Курсы не найдены.", size=12, italic=True)

    doc.add_paragraph()

    # Итог
    completed = [c for c in courses if c.get('status') == 'completed']
    grades    = [c['grade'] for c in completed if c.get('grade') is not None]
    avg_grade = round(sum(grades) / len(grades), 1) if grades else None

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(1.25)
    _add_run(p,
             f"Всего курсов: {len(courses)}. Завершено: {len(completed)}. "
             + (f"Средний балл: {avg_grade}." if avg_grade else ""))

    doc.add_paragraph()
    _para(doc, "Справка выдана по месту требования.", size=12, italic=True,
          align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    _add_line(doc)
    doc.add_paragraph()

    _signature_row(doc,
                   "Начальник учебного отдела",
                   f"{now.strftime('%d.%m.%Y')}")
    doc.add_paragraph()
    _signature_row(doc,
                   "_________________ / И.О. Фамилия /",
                   "М.П.")
    doc.add_paragraph()
    _para(doc,
          "Ректор: ____________________  / И.О. Фамилия /",
          size=11, space_before=4)

    return _to_buffer(doc)


# ─── 3. Сертификат об окончании курса ────────────────────────────────────────

def generate_certificate(student_name: str, course_title: str,
                         teacher_name: str, grade: float,
                         completion_date: str = None) -> bytes:
    """Сертификат об успешном прохождении курса"""
    doc = Document()
    _set_page_margins(doc, top=1.5, bottom=1.5, left=2, right=2)
    now = datetime.now()
    months_ru = {
        'January':'января','February':'февраля','March':'марта','April':'апреля',
        'May':'мая','June':'июня','July':'июля','August':'августа',
        'September':'сентября','October':'октября','November':'ноября','December':'декабря'
    }

    date_str = completion_date or (
            now.strftime('%d') + ' ' + months_ru[now.strftime('%B')] + ' ' + now.strftime('%Y') + ' г.'
    )

    # Рамка через таблицу 1×1
    outer = doc.add_table(rows=1, cols=1)
    outer.style = 'Table Grid'
    cell = outer.rows[0].cells[0]

    # Толстая цветная рамка
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top','bottom','left','right'):
        bd = OxmlElement(f'w:{side}')
        bd.set(qn('w:val'),   'single')
        bd.set(qn('w:sz'),    '18')
        bd.set(qn('w:color'), '4B3F8C')
        tcBorders.append(bd)
    tcPr.append(tcBorders)

    # Содержимое сертификата внутри ячейки
    def cp(text="", align=WD_ALIGN_PARAGRAPH.CENTER, size=12, bold=False,
           italic=False, space_before=0, space_after=8, color=None):
        p = cell.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after  = Pt(space_after)
        if text:
            run = p.add_run(text)
            _set_font(run, "Times New Roman", size, bold, italic, color)
        return p

    cell.paragraphs[0].clear()  # убрать пустой первый параграф

    cp(space_before=16, space_after=4)
    cp("★  УНИВЕРСИТЕТ  ★",
       size=13, bold=True, space_after=4, color=(75, 63, 140))

    cp("СЕРТИФИКАТ",
       size=26, bold=True, space_before=8, space_after=2, color=(75, 63, 140))
    cp("об успешном прохождении курса",
       size=14, italic=True, space_after=20, color=(100, 90, 160))

    # Разделитель
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(16)
    run = p.add_run("─" * 48)
    _set_font(run, "Times New Roman", 11, color=(75, 63, 140))

    cp("Настоящим подтверждается, что", size=13, space_after=4)
    cp(student_name, size=22, bold=True, space_after=4, color=(30, 30, 30))
    cp("успешно прошёл(-а) обучение по курсу", size=13, space_after=4)
    cp(f"«{course_title}»", size=18, bold=True, space_after=16, color=(75, 63, 140))

    # Оценка
    if grade is not None:
        if grade >= 90:
            level = "с отличием"
            level_color = (22, 163, 74)
        elif grade >= 75:
            level = "с хорошим результатом"
            level_color = (2, 132, 199)
        elif grade >= 60:
            level = "с удовлетворительным результатом"
            level_color = (217, 119, 6)
        else:
            level = ""
            level_color = (100, 100, 100)

        cp(f"Итоговая оценка: {grade} баллов", size=14, bold=True, space_after=4)
        if level:
            cp(f"Курс пройден {level}", size=13, italic=True,
               space_after=16, color=level_color)

    # Разделитель
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(16)
    run = p.add_run("─" * 48)
    _set_font(run, "Times New Roman", 11, color=(75, 63, 140))

    cp(f"Дата выдачи: {date_str}", size=12, space_after=4)
    if teacher_name:
        cp(f"Преподаватель: {teacher_name}", size=12, space_after=20)

    # Подписи
    tbl = cell.add_table(rows=1, cols=3)
    sig_cols = [Cm(5), Cm(4), Cm(5)]
    sig_texts = [
        f"Студент:\n\n_______________\n{student_name}",
        "М.П.",
        f"Ректор:\n\n_______________\n/ И.О. Фамилия /",
    ]
    sig_aligns = [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.RIGHT]

    for cell_s, text, align, w in zip(tbl.rows[0].cells, sig_texts, sig_aligns, sig_cols):
        # Убрать границы
        tc2 = cell_s._tc
        tcPr2 = tc2.get_or_add_tcPr()
        tcBorders2 = OxmlElement('w:tcBorders')
        for side in ('top','bottom','left','right'):
            bd = OxmlElement(f'w:{side}')
            bd.set(qn('w:val'), 'nil')
            tcBorders2.append(bd)
        tcPr2.append(tcBorders2)
        cell_s.width = w

        p = cell_s.paragraphs[0]
        p.alignment = align
        run = p.add_run(text)
        _set_font(run, "Times New Roman", 10)

    cell.add_paragraph().paragraph_format.space_after = Pt(16)

    return _to_buffer(doc)